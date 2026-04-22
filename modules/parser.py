from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader


def _clean_text(value) -> str:
    if value is None:
        return ""
    return " ".join(str(value).replace("\xa0", " ").split()).strip()


def extract_docx_text(file_path: str | Path) -> str:
    file_path = Path(file_path)
    doc = Document(str(file_path))

    parts: list[str] = []

    # Paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = [_clean_text(cell.text) for cell in row.cells]
            row_cells = [cell for cell in row_cells if cell]
            if row_cells:
                parts.append(" | ".join(row_cells))

    return "\n".join(parts).strip()


def extract_docx_tables(file_path: str | Path) -> list[dict]:
    file_path = Path(file_path)
    doc = Document(str(file_path))

    extracted_tables: list[dict] = []

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []

        for row in table.rows:
            row_cells = [_clean_text(cell.text) for cell in row.cells]
            rows.append(row_cells)

        extracted_tables.append({
            "table_number": table_index,
            "rows": rows,
        })

    return extracted_tables


def extract_pdf_text(file_path: str | Path) -> str:
    file_path = Path(file_path)
    reader = PdfReader(str(file_path))

    pages_text: list[str] = []

    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        text = text.replace("\xa0", " ").strip()
        if text:
            pages_text.append(text)

    return "\n\n".join(pages_text).strip()


def extract_pdf_tables(file_path: str | Path) -> list[dict]:
    """
    For now we support PDF text extraction only.
    Table extraction from PDFs can be added later.
    """
    return []


def extract_file_text(file_path: str | Path) -> str:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return extract_docx_text(file_path)

    if suffix == ".pdf":
        return extract_pdf_text(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")


def extract_file_tables(file_path: str | Path) -> list[dict]:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return extract_docx_tables(file_path)

    if suffix == ".pdf":
        return extract_pdf_tables(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")