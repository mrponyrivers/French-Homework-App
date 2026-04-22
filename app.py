
from __future__ import annotations

import html
import io
import random
import re
import sqlite3
import unicodedata
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
from docx import Document

# -------------------------------------------------
# Page config
# -------------------------------------------------
st.set_page_config(
    page_title="French Homework App",
    page_icon="🇫🇷",
    layout="wide",
)

APP_TITLE = "🇫🇷 French Homework App"
APP_SUBTITLE = (
    "Upload French homework files and turn them into saved lessons, "
    "vocabulary, flashcards, quizzes, and exercise practice."
)

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)

# Clean app-specific DB to avoid legacy schema collisions.
DB_PATH = DATA_DIR / "french_homework_app_v2.db"

FRENCH_HEADER_HINTS = {
    "french", "français", "francais", "fr", "mot", "mot français", "word", "terme"
}
ENGLISH_HEADER_HINTS = {
    "english", "anglais", "eng", "meaning", "translation", "traduction", "definition"
}
BLANK_PATTERN = re.compile(r"(_{1,}|\.{3,}|\[\s*\]|\(\s*\))")
ANSWER_AT_END_PATTERN = re.compile(r"^(.*?)(?:\s*[\(\[]\s*([^()\[\]\n]{1,80})\s*[\)\]])\s*$")
BLANK_SEGMENT_PATTERN = re.compile(
    r"((?:_{1,}|\.{3,}|\[\s*\]|\(\s*\)).*?)\s*\(([^()\n]{1,80})\)"
)
EXERCISE_HEADER_PATTERN = re.compile(r"^(exercice|exercise)\s*([0-9]+)?[\.:\-]?\s*(.*)$", re.IGNORECASE)
INSTRUCTION_HINT_PATTERN = re.compile(
    r"\b(complete|compl[eé]te|choisis|choisir|avec|fill|conjugue|conjuguez|write|rewrite|answer|transform)\b",
    re.IGNORECASE,
)

# -------------------------------------------------
# Database
# -------------------------------------------------
def get_conn() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_storage() -> None:
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS lessons (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            source_filename TEXT,
            raw_text TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS exercises (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lesson_id INTEGER NOT NULL,
            exercise_title TEXT,
            prompt TEXT NOT NULL,
            answer TEXT,
            exercise_type TEXT,
            position INTEGER,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (lesson_id) REFERENCES lessons (id) ON DELETE CASCADE
        )
        """
    )

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS vocab_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lesson_id INTEGER NOT NULL,
            term TEXT NOT NULL,
            meaning TEXT,
            source TEXT,
            position INTEGER,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (lesson_id) REFERENCES lessons (id) ON DELETE CASCADE
        )
        """
    )

    conn.commit()
    conn.close()


# -------------------------------------------------
# Helpers
# -------------------------------------------------
def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def clean_multiline_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ")
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def strip_accents(value: Any) -> str:
    text = clean_text(value).replace("œ", "oe").replace("Œ", "OE").replace("æ", "ae").replace("Æ", "AE")
    text = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in text if not unicodedata.combining(ch))


def strip_bold_markers(text: str) -> str:
    text = clean_multiline_text(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    return text


def extract_bold_instruction_text(text: str) -> str:
    text = clean_multiline_text(text)
    parts: List[str] = []

    for match in re.findall(r"\*\*(.*?)\*\*", text):
        cleaned = clean_multiline_text(match)
        if cleaned:
            parts.append(f"**{cleaned}**")

    for match in re.findall(r"__(.*?)__", text):
        cleaned = clean_multiline_text(match)
        if cleaned:
            parts.append(f"**{cleaned}**")

    return "\n".join(parts)


def extract_instruction_display_text(line: str) -> str:
    line = clean_multiline_text(line)
    bold_only = extract_bold_instruction_text(line)
    if bold_only:
        return bold_only

    stripped = strip_bold_markers(line)
    header_match = EXERCISE_HEADER_PATTERN.match(stripped)
    if header_match:
        trailing = clean_multiline_text(header_match.group(3))
        if trailing:
            return trailing

    return stripped


def normalize_answer_for_compare(value: Any) -> str:
    return normalize_for_compare(strip_accents(value))


def format_runs_as_marked_text(runs) -> str:
    parts: List[str] = []
    for run in runs:
        run_text = run.text.replace("\xa0", " ")
        if not run_text:
            continue
        if getattr(run, "bold", False):
            parts.append(f"**{run_text}**")
        else:
            parts.append(run_text)

    text = "".join(parts)
    text = re.sub(r"[ 	]+", " ", text)
    return clean_multiline_text(text)


def normalize_for_compare(value: Any) -> str:
    text = clean_text(value).casefold()
    text = text.replace("’", "'")
    text = re.sub(r"[^a-z0-9àâçéèêëîïôûùüÿñæœ'\- ]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def unique_preserve_order(items: List[Dict[str, Any]], key_fields: List[str]) -> List[Dict[str, Any]]:
    seen = set()
    result: List[Dict[str, Any]] = []
    for item in items:
        key = tuple(normalize_for_compare(item.get(field, "")) for field in key_fields)
        if key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


def lesson_display_number_map(lessons: List[Dict[str, Any]]) -> Dict[int, int]:
    # lessons are fetched newest first. Make newest show the highest lesson count.
    total = len(lessons)
    mapping: Dict[int, int] = {}
    for idx, lesson in enumerate(lessons):
        mapping[int(lesson["id"])] = total - idx
    return mapping


def row_to_lesson_dict(row: sqlite3.Row) -> Dict[str, Any]:
    return {
        "id": int(row["id"]),
        "title": clean_text(row["title"]),
        "source_filename": clean_text(row["source_filename"]),
        "raw_text": clean_multiline_text(row["raw_text"]),
        "created_at": clean_text(row["created_at"]),
    }


def row_to_exercise_dict(row: sqlite3.Row) -> Dict[str, Any]:
    return {
        "id": int(row["id"]),
        "lesson_id": int(row["lesson_id"]),
        "exercise_title": clean_text(row["exercise_title"]),
        "prompt": clean_multiline_text(row["prompt"]),
        "answer": clean_text(row["answer"]),
        "exercise_type": clean_text(row["exercise_type"] or "general"),
        "position": int(row["position"] or 0),
    }


def row_to_vocab_dict(row: sqlite3.Row) -> Dict[str, Any]:
    return {
        "id": int(row["id"]),
        "lesson_id": int(row["lesson_id"]),
        "term": clean_text(row["term"]),
        "meaning": clean_text(row["meaning"]),
        "source": clean_text(row["source"] or "table"),
        "position": int(row["position"] or 0),
    }


# -------------------------------------------------
# DB reads/writes
# -------------------------------------------------
def fetch_lessons() -> List[Dict[str, Any]]:
    conn = get_conn()
    rows = conn.execute(
        """
        SELECT id, title, source_filename, raw_text, created_at
        FROM lessons
        ORDER BY id DESC
        """
    ).fetchall()
    conn.close()
    return [row_to_lesson_dict(row) for row in rows]


def fetch_lesson_by_id(lesson_id: int) -> Optional[Dict[str, Any]]:
    conn = get_conn()
    row = conn.execute(
        """
        SELECT id, title, source_filename, raw_text, created_at
        FROM lessons
        WHERE id = ?
        """,
        (lesson_id,),
    ).fetchone()
    conn.close()
    return row_to_lesson_dict(row) if row else None


def fetch_exercises_for_lesson(lesson_id: int) -> List[Dict[str, Any]]:
    conn = get_conn()
    rows = conn.execute(
        """
        SELECT id, lesson_id, exercise_title, prompt, answer, exercise_type, position
        FROM exercises
        WHERE lesson_id = ?
        ORDER BY position, id
        """,
        (lesson_id,),
    ).fetchall()
    conn.close()
    return [row_to_exercise_dict(row) for row in rows]


def fetch_vocab_for_lesson(lesson_id: int) -> List[Dict[str, Any]]:
    conn = get_conn()
    rows = conn.execute(
        """
        SELECT id, lesson_id, term, meaning, source, position
        FROM vocab_items
        WHERE lesson_id = ?
        ORDER BY position, id
        """,
        (lesson_id,),
    ).fetchall()
    conn.close()
    return [row_to_vocab_dict(row) for row in rows]


def save_lesson_record(title: str, source_filename: str, raw_text: str) -> Tuple[int, int]:
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO lessons (title, source_filename, raw_text)
        VALUES (?, ?, ?)
        """,
        (clean_text(title) or "Untitled Lesson", clean_text(source_filename), clean_multiline_text(raw_text)),
    )
    lesson_id = int(cur.lastrowid)
    conn.commit()
    conn.close()

    total_count = len(fetch_lessons())
    return lesson_id, total_count


def save_exercise_records(lesson_id: int, exercises: List[Dict[str, Any]]) -> None:
    clean_items: List[Dict[str, Any]] = []

    for idx, item in enumerate(exercises, start=1):
        prompt = clean_multiline_text(item.get("prompt", ""))
        if not prompt:
            continue

        answer = clean_text(item.get("answer", ""))
        exercise_type = clean_text(item.get("exercise_type", "general")) or "general"
        position = int(item.get("position") or idx)

        clean_items.append(
            {
                "lesson_id": lesson_id,
                "exercise_title": clean_text(item.get("exercise_title")) or f"Exercise {position}",
                "prompt": prompt,
                "answer": answer,
                "exercise_type": exercise_type,
                "position": position,
            }
        )

    clean_items = unique_preserve_order(clean_items, ["prompt", "answer"])

    if not clean_items:
        return

    conn = get_conn()
    conn.executemany(
        """
        INSERT INTO exercises (lesson_id, exercise_title, prompt, answer, exercise_type, position)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        [
            (
                item["lesson_id"],
                item["exercise_title"],
                item["prompt"],
                item["answer"],
                item["exercise_type"],
                item["position"],
            )
            for item in clean_items
        ],
    )
    conn.commit()
    conn.close()


def save_vocab_records(lesson_id: int, vocab_items: List[Dict[str, Any]]) -> None:
    clean_items: List[Dict[str, Any]] = []

    for idx, item in enumerate(vocab_items, start=1):
        term = clean_text(item.get("term", ""))
        meaning = clean_text(item.get("meaning", ""))
        if not term:
            continue
        clean_items.append(
            {
                "lesson_id": lesson_id,
                "term": term,
                "meaning": meaning,
                "source": clean_text(item.get("source", "table")) or "table",
                "position": int(item.get("position") or idx),
            }
        )

    clean_items = unique_preserve_order(clean_items, ["term", "meaning"])

    if not clean_items:
        return

    conn = get_conn()
    conn.executemany(
        """
        INSERT INTO vocab_items (lesson_id, term, meaning, source, position)
        VALUES (?, ?, ?, ?, ?)
        """,
        [
            (item["lesson_id"], item["term"], item["meaning"], item["source"], item["position"])
            for item in clean_items
        ],
    )
    conn.commit()
    conn.close()


def delete_lesson_record(lesson_id: int) -> None:
    conn = get_conn()
    conn.execute("DELETE FROM lessons WHERE id = ?", (lesson_id,))
    conn.commit()
    conn.close()


def delete_last_n_lessons(n: int = 2) -> int:
    lessons = fetch_lessons()
    ids_to_delete = [lesson["id"] for lesson in lessons[:n]]
    for lesson_id in ids_to_delete:
        delete_lesson_record(int(lesson_id))
    return len(ids_to_delete)


# -------------------------------------------------
# DOCX parsing
# -------------------------------------------------
def extract_docx_content(uploaded_file) -> Tuple[str, List[List[List[str]]]]:
    file_bytes = uploaded_file.getvalue()
    doc = Document(io.BytesIO(file_bytes))

    paragraphs: List[str] = []
    for para in doc.paragraphs:
        text = format_runs_as_marked_text(para.runs) if para.runs else clean_multiline_text(para.text)
        if text:
            paragraphs.append(text)

    tables: List[List[List[str]]] = []
    for table in doc.tables:
        table_rows: List[List[str]] = []
        for row in table.rows:
            row_cells: List[str] = []
            for cell in row.cells:
                cell_lines: List[str] = []
                for para in cell.paragraphs:
                    para_text = format_runs_as_marked_text(para.runs) if para.runs else clean_multiline_text(para.text)
                    if para_text:
                        cell_lines.append(para_text)
                cell_text = "\n".join(cell_lines) if cell_lines else clean_multiline_text(cell.text)
                row_cells.append(cell_text)
            if any(cell for cell in row_cells):
                table_rows.append(row_cells)
        if table_rows:
            tables.append(table_rows)

    raw_text = "\n".join(paragraphs)
    return raw_text, tables


def is_vocab_header_row(row: List[str]) -> Tuple[bool, Optional[int], Optional[int]]:
    normalized = [normalize_for_compare(cell) for cell in row]
    french_idx = None
    english_idx = None

    for idx, cell in enumerate(normalized):
        if any(hint in cell for hint in FRENCH_HEADER_HINTS):
            french_idx = idx
        if any(hint in cell for hint in ENGLISH_HEADER_HINTS):
            english_idx = idx

    return french_idx is not None and english_idx is not None, french_idx, english_idx


def looks_like_header_cell(cell: str) -> bool:
    norm = normalize_for_compare(cell)
    return (
        norm in FRENCH_HEADER_HINTS
        or norm in ENGLISH_HEADER_HINTS
        or "french" in norm
        or "english" in norm
        or norm in {"rank", "number", "no", "#"}
    )


def extract_vocab_pairs_from_tables(tables: List[List[List[str]]]) -> List[Dict[str, Any]]:
    vocab_items: List[Dict[str, Any]] = []
    position = 1

    for table in tables:
        if not table:
            continue

        french_col = None
        english_col = None
        start_row = 0

        for row_idx, row in enumerate(table[:3]):
            is_header, detected_french_col, detected_english_col = is_vocab_header_row(row)
            if is_header:
                french_col = detected_french_col
                english_col = detected_english_col
                start_row = row_idx + 1
                break

        if french_col is None or english_col is None:
            max_cols = max(len(row) for row in table)
            if max_cols >= 3:
                french_col = 1
                english_col = 2
                start_row = 1 if any(looks_like_header_cell(c) for c in table[0]) else 0
            elif max_cols >= 2:
                french_col = 0
                english_col = 1
                start_row = 1 if any(looks_like_header_cell(c) for c in table[0]) else 0
            else:
                continue

        for row in table[start_row:]:
            if len(row) <= max(french_col, english_col):
                continue
            term = clean_text(row[french_col])
            meaning = clean_text(row[english_col])

            if not term:
                continue
            if looks_like_header_cell(term) or looks_like_header_cell(meaning):
                continue
            if re.fullmatch(r"\d+", term):
                continue

            vocab_items.append(
                {
                    "term": term,
                    "meaning": meaning,
                    "source": "table",
                    "position": position,
                }
            )
            position += 1

    return unique_preserve_order(vocab_items, ["term", "meaning"])


# -------------------------------------------------
# Exercise parsing
# -------------------------------------------------
def has_blank_pattern(text: str) -> bool:
    return bool(BLANK_PATTERN.search(text))


def parse_fill_blank_prompt_answer(text: str) -> Tuple[str, str]:
    text = clean_multiline_text(text)
    match = ANSWER_AT_END_PATTERN.match(text)
    if match:
        prompt = clean_multiline_text(match.group(1))
        answer = clean_text(match.group(2))
        if has_blank_pattern(prompt) and answer and not has_blank_pattern(answer):
            return prompt, answer
    return text, ""


def normalize_blank_prompt(prompt: str) -> str:
    prompt = clean_multiline_text(prompt)
    prompt = re.sub(r"^([_\.\[\(]\s*)", "_ ", prompt)
    prompt = re.sub(r"^_([A-Za-zÀ-ÿ])", r"_ \1", prompt)
    prompt = re.sub(r"\s+", " ", prompt).strip()
    return prompt


def extract_exercise_title(line: str) -> str:
    line = strip_bold_markers(line)
    match = EXERCISE_HEADER_PATTERN.match(line)
    if not match:
        return ""
    number = clean_text(match.group(2))
    return f"Exercise {number}" if number else "Exercise"


def is_section_instruction_line(line: str) -> bool:
    raw_line = clean_multiline_text(line)
    line = strip_bold_markers(raw_line)
    if not line or has_blank_pattern(line):
        return False
    return bool(
        EXERCISE_HEADER_PATTERN.match(line)
        or INSTRUCTION_HINT_PATTERN.search(line)
        or extract_bold_instruction_text(raw_line)
    )


def combine_instruction_with_prompt(instruction: str, prompt: str) -> str:
    instruction = clean_multiline_text(instruction)
    prompt = normalize_blank_prompt(prompt)
    if instruction:
        return f"{instruction}\n{prompt}"

    return prompt


def extract_blank_lines_from_text(raw_text: str) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    position = 1

    lines = [clean_multiline_text(line) for line in raw_text.splitlines()]
    lines = [line for line in lines if line]

    current_instruction_lines: List[str] = []
    current_exercise_title = ""

    for line in lines:
        stripped_line = strip_bold_markers(line)

        if is_section_instruction_line(line):
            extracted_title = extract_exercise_title(stripped_line)
            if extracted_title:
                current_exercise_title = extracted_title
                current_instruction_lines = []

            instruction_text = extract_instruction_display_text(line)
            if instruction_text and instruction_text not in current_instruction_lines:
                current_instruction_lines.append(instruction_text)
            continue

        if not has_blank_pattern(line):
            continue

        prompt, answer = parse_fill_blank_prompt_answer(line)
        instruction_block = "\n".join(current_instruction_lines)

        prompt = combine_instruction_with_prompt(instruction_block, prompt)

        if has_blank_pattern(prompt):
            items.append(
                {
                    "exercise_title": current_exercise_title or f"Exercise {position}",
                    "prompt": prompt,
                    "answer": answer,
                    "exercise_type": "fill_blank",
                    "position": position,
                }
            )
            position += 1

    if items:
        return unique_preserve_order(items, ["prompt", "answer"])

    # Fallback: handle dense single-line text where multiple blank questions collapsed together.
    merged_text = clean_multiline_text(raw_text)
    current_exercise_title = ""
    title_match = EXERCISE_HEADER_PATTERN.search(merged_text)
    if title_match:
        number = clean_text(title_match.group(2))
        current_exercise_title = f"Exercise {number}" if number else "Exercise"

    for match in BLANK_SEGMENT_PATTERN.finditer(merged_text):
        prompt = normalize_blank_prompt(match.group(1))
        answer = clean_text(match.group(2))
        if has_blank_pattern(prompt):
            items.append(
                {
                    "exercise_title": current_exercise_title or f"Exercise {position}",
                    "prompt": prompt,
                    "answer": answer,
                    "exercise_type": "fill_blank",
                    "position": position,
                }
            )
            position += 1

    return unique_preserve_order(items, ["prompt", "answer"])


def split_into_exercises(raw_text: str) -> List[Dict[str, Any]]:
    blank_items = extract_blank_lines_from_text(raw_text)
    if blank_items:
        return blank_items

    lines = [clean_multiline_text(line) for line in raw_text.splitlines()]
    lines = [line for line in lines if line]
    if not lines:
        return []

    exercise_blocks: List[str] = []
    buffer: List[str] = []
    start_pattern = re.compile(r"^(exercice|exercise|question|q[0-9]*|[0-9]+[.)\-:])\s*", re.IGNORECASE)

    for line in lines:
        if start_pattern.match(line):
            if buffer:
                exercise_blocks.append("\n".join(buffer))
            buffer = [line]
        else:
            if buffer:
                buffer.append(line)

    if buffer:
        exercise_blocks.append("\n".join(buffer))

    if not exercise_blocks and raw_text.strip():
        paragraphs = [p.strip() for p in raw_text.split("\n\n") if p.strip()]
        exercise_blocks = paragraphs[:12]

    exercises: List[Dict[str, Any]] = []
    for idx, block in enumerate(exercise_blocks, start=1):
        exercises.append(
            {
                "exercise_title": f"Exercise {idx}",
                "prompt": clean_multiline_text(block),
                "answer": "",
                "exercise_type": "general",
                "position": idx,
            }
        )

    return exercises


def guess_answer_from_prompt(prompt: str) -> str:
    prompt = clean_multiline_text(prompt)

    match = re.search(r"\(([^()]{1,40})\)\s*$", prompt)
    if match:
        candidate = clean_text(match.group(1))
        if candidate and not has_blank_pattern(candidate):
            return candidate

    match = re.search(r"\[([^\[\]]{1,40})\]\s*$", prompt)
    if match:
        candidate = clean_text(match.group(1))
        if candidate and not has_blank_pattern(candidate):
            return candidate

    return ""


def build_fill_blank_items(exercises: List[Dict[str, Any]], raw_text: str) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []

    for exercise in exercises:
        prompt = normalize_blank_prompt(exercise.get("prompt", ""))
        answer = clean_text(exercise.get("answer", ""))
        exercise_type = clean_text(exercise.get("exercise_type", ""))

        if not prompt or not has_blank_pattern(prompt):
            continue

        items.append(
            {
                "prompt": prompt,
                "answer": answer or guess_answer_from_prompt(prompt),
                "position": int(exercise.get("position") or len(items) + 1),
            }
        )

    if items:
        return unique_preserve_order(items, ["prompt", "answer"])

    raw_items = extract_blank_lines_from_text(raw_text)
    return [
        {"prompt": item["prompt"], "answer": item["answer"], "position": item["position"]}
        for item in raw_items
    ]


def build_exercise_quiz_items(exercises: List[Dict[str, Any]], raw_text: str) -> List[Dict[str, Any]]:
    base_items = build_fill_blank_items(exercises, raw_text)
    return [item for item in base_items if clean_text(item.get("answer", ""))]


def build_vocab_hint_pool(vocab_items: List[Dict[str, Any]], limit: int = 8) -> List[Dict[str, Any]]:
    if not vocab_items:
        return []
    return sorted(vocab_items, key=lambda x: normalize_for_compare(x["term"]))[:limit]


def render_big_card(text: str, answer: bool = False) -> None:
    safe_text = html.escape(clean_multiline_text(text))
    safe_text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", safe_text)
    safe_text = re.sub(r"__(.+?)__", r"<strong>\1</strong>", safe_text)
    safe_text = safe_text.replace("\n", "<br>")
    extra = " answer" if answer else ""
    st.markdown(f'<div class="big-card{extra}">{safe_text}</div>', unsafe_allow_html=True)


def clear_answer_inputs() -> None:
    st.session_state["quiz_input"] = ""
    st.session_state["practice_input"] = ""
    st.session_state["exercise_quiz_input"] = ""
    st.session_state["quiz_input_nonce"] = st.session_state.get("quiz_input_nonce", 0) + 1
    st.session_state["practice_input_nonce"] = st.session_state.get("practice_input_nonce", 0) + 1
    st.session_state["exercise_quiz_input_nonce"] = st.session_state.get("exercise_quiz_input_nonce", 0) + 1


def inject_custom_css() -> None:
    st.markdown(
        """
        <style>
        .big-card {
            font-size: 1.42rem !important;
            line-height: 1.34 !important;
            font-weight: 700 !important;
            padding: 1.25rem 1.35rem !important;
            border-radius: 18px !important;
            border: 1px solid rgba(120, 120, 140, 0.22) !important;
            background: #ffffff !important;
            color: #111827 !important;
            box-shadow: 0 6px 18px rgba(15, 23, 42, 0.06) !important;
            margin: 0.25rem 0 1rem 0 !important;
        }
        .big-card.answer {
            font-size: 1.2rem !important;
            background: #ecfdf5 !important;
            color: #14532d !important;
        }
        .section-label {
            font-size: 0.95rem;
            font-weight: 700;
            letter-spacing: 0.03em;
            text-transform: uppercase;
            color: #475569;
            margin-top: 0.35rem;
            margin-bottom: 0.45rem;
        }
        [data-testid="stMetric"] {
            background: #f3f4f6 !important;
            border: 1px solid rgba(148, 163, 184, 0.28) !important;
            border-radius: 16px !important;
            padding: 0.65rem 0.8rem !important;
        }
        [data-testid="stMetricLabel"],
        [data-testid="stMetricLabel"] *,
        [data-testid="stMetricValue"],
        [data-testid="stMetricValue"] *,
        [data-testid="stMetricDelta"],
        [data-testid="stMetricDelta"] * {
            color: #111827 !important;
            fill: #111827 !important;
        }
        [data-testid="stMetricLabel"] {
            font-size: 0.82rem;
            font-weight: 600 !important;
        }
        [data-testid="stMetricValue"] {
            font-size: 1.35rem;
            font-weight: 700 !important;
        }
        .stButton > button {
            border-radius: 12px;
            font-weight: 600;
        }
        .stTextInput input {
            font-size: 1.05rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# -------------------------------------------------
# Flashcard and quiz helpers
# -------------------------------------------------
def build_flashcard_deck(vocab_items: List[Dict[str, Any]], order_mode: str, start_letter: str = "") -> List[Dict[str, Any]]:
    deck = [dict(item) for item in vocab_items]
    if not deck:
        return deck

    if order_mode == "A → Z":
        deck = sorted(deck, key=lambda x: normalize_for_compare(x["term"]))
    elif order_mode == "Random starting point":
        deck = sorted(deck, key=lambda x: normalize_for_compare(x["term"]))
        start_idx = random.randint(0, len(deck) - 1)
        deck = deck[start_idx:] + deck[:start_idx]
    elif order_mode == "Full shuffle":
        random.shuffle(deck)
    elif order_mode == "Start from letter":
        deck = sorted(deck, key=lambda x: normalize_for_compare(x["term"]))
        letter = clean_text(start_letter[:1]).upper()
        if letter:
            matching_indexes = [
                idx for idx, item in enumerate(deck)
                if clean_text(item["term"][:1]).upper() == letter
            ]
            if matching_indexes:
                start_idx = matching_indexes[0]
                deck = deck[start_idx:] + deck[:start_idx]
    return deck


def flashcard_signature(lesson_id: int, order_mode: str, start_letter: str, vocab_items: List[Dict[str, Any]]) -> str:
    base = ",".join(f"{item.get('id', '')}:{item['term']}:{item['meaning']}" for item in vocab_items)
    return f"{lesson_id}|{order_mode}|{start_letter.upper()}|{base}"


def answer_variants(expected: str) -> List[str]:
    expected = clean_text(expected)
    if not expected:
        return []

    variants = {normalize_answer_for_compare(expected)}
    parts = re.split(r"/|;|,|\bor\b", expected, flags=re.IGNORECASE)

    for part in parts:
        cleaned = normalize_answer_for_compare(part)
        if cleaned:
            variants.add(cleaned)

    for part in parts:
        stripped = re.sub(r"\([^)]*\)", "", part).strip()
        cleaned = normalize_answer_for_compare(stripped)
        if cleaned:
            variants.add(cleaned)

    return [variant for variant in variants if variant]


def is_answer_correct(user_answer: str, expected_answer: str) -> bool:
    user_forms = {
        normalize_answer_for_compare(user_answer),
        normalize_for_compare(clean_text(user_answer)),
    }
    user_forms = {form for form in user_forms if form}
    if not user_forms:
        return False

    for variant in answer_variants(expected_answer):
        for user_norm in user_forms:
            if user_norm == variant:
                return True
            if len(user_norm) >= 3 and user_norm in variant:
                return True
            if len(variant) >= 3 and variant in user_norm:
                return True

    return False


def quiz_signature(lesson_id: int, direction: str, vocab_items: List[Dict[str, Any]]) -> str:
    base = ",".join(f"{item.get('id', '')}:{item['term']}:{item['meaning']}" for item in vocab_items)
    return f"{lesson_id}|{direction}|{base}"


def exercise_quiz_signature(lesson_id: int, items: List[Dict[str, Any]]) -> str:
    base = ",".join(f"{item['position']}:{item['prompt']}:{item['answer']}" for item in items)
    return f"{lesson_id}|{base}"


def practice_signature(lesson_id: int, items: List[Dict[str, Any]]) -> str:
    base = ",".join(f"{item['position']}:{item['prompt']}:{item['answer']}" for item in items)
    return f"{lesson_id}|{base}"


# -------------------------------------------------
# Session state
# -------------------------------------------------
def ensure_session_state() -> None:
    defaults = {
        "flashcard_deck": [],
        "flashcard_index": 0,
        "flashcard_show_back": False,
        "flashcard_signature": "",
        "quiz_deck": [],
        "quiz_index": 0,
        "quiz_score": 0,
        "quiz_feedback": "",
        "quiz_answered": False,
        "quiz_signature": "",
        "quiz_input": "",
        "quiz_input_nonce": 0,
        "practice_deck": [],
        "practice_index": 0,
        "practice_feedback": "",
        "practice_answered": False,
        "practice_signature": "",
        "practice_input": "",
        "practice_input_nonce": 0,
        "exercise_quiz_deck": [],
        "exercise_quiz_index": 0,
        "exercise_quiz_score": 0,
        "exercise_quiz_feedback": "",
        "exercise_quiz_answered": False,
        "exercise_quiz_signature": "",
        "exercise_quiz_input": "",
        "exercise_quiz_input_nonce": 0,
        "last_save_message": "",
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def reset_flashcards(lesson_id: int, vocab_items: List[Dict[str, Any]], order_mode: str, start_letter: str) -> None:
    st.session_state.flashcard_deck = build_flashcard_deck(vocab_items, order_mode, start_letter)
    st.session_state.flashcard_index = 0
    st.session_state.flashcard_show_back = False
    st.session_state.flashcard_signature = flashcard_signature(lesson_id, order_mode, start_letter, vocab_items)


def reset_quiz(lesson_id: int, vocab_items: List[Dict[str, Any]], direction: str) -> None:
    deck = [dict(item) for item in vocab_items]
    random.shuffle(deck)

    st.session_state.quiz_deck = deck
    st.session_state.quiz_index = 0
    st.session_state.quiz_score = 0
    st.session_state.quiz_feedback = ""
    st.session_state.quiz_answered = False
    st.session_state.quiz_input = ""
    st.session_state.quiz_input_nonce = st.session_state.get("quiz_input_nonce", 0) + 1
    st.session_state.quiz_signature = quiz_signature(lesson_id, direction, vocab_items)


def reset_practice(lesson_id: int, items: List[Dict[str, Any]]) -> None:
    st.session_state.practice_deck = [dict(item) for item in items]
    st.session_state.practice_index = 0
    st.session_state.practice_feedback = ""
    st.session_state.practice_answered = False
    st.session_state.practice_input = ""
    st.session_state.practice_input_nonce = st.session_state.get("practice_input_nonce", 0) + 1
    st.session_state.practice_signature = practice_signature(lesson_id, items)


def reset_exercise_quiz(lesson_id: int, items: List[Dict[str, Any]]) -> None:
    deck = [dict(item) for item in items]
    random.shuffle(deck)

    st.session_state.exercise_quiz_deck = deck
    st.session_state.exercise_quiz_index = 0
    st.session_state.exercise_quiz_score = 0
    st.session_state.exercise_quiz_feedback = ""
    st.session_state.exercise_quiz_answered = False
    st.session_state.exercise_quiz_input = ""
    st.session_state.exercise_quiz_input_nonce = st.session_state.get("exercise_quiz_input_nonce", 0) + 1
    st.session_state.exercise_quiz_signature = exercise_quiz_signature(lesson_id, items)


# -------------------------------------------------
# App init
# -------------------------------------------------
init_storage()
ensure_session_state()
inject_custom_css()

st.title(APP_TITLE)
st.caption(APP_SUBTITLE)

tabs = st.tabs(
    [
        "Upload Lesson",
        "Lesson History",
        "Flashcards",
        "Quiz Mode",
        "Exercise Practice",
    ]
)

lessons = fetch_lessons()
display_numbers = lesson_display_number_map(lessons)
lesson_map = {f"Lesson {display_numbers[lesson['id']]} — {lesson['title']}": lesson["id"] for lesson in lessons}


# -------------------------------------------------
# Tab 1: Upload Lesson
# -------------------------------------------------
with tabs[0]:
    st.subheader("Upload a new French homework file")

    if st.session_state["last_save_message"]:
        st.success(st.session_state["last_save_message"])
        st.session_state["last_save_message"] = ""

    st.caption("This version uses a clean app database so lesson numbers stay 1, 2, 3 instead of inheriting old broken IDs.")

    uploaded_file = st.file_uploader(
        "Upload a .docx lesson file",
        type=["docx"],
        key="lesson_uploader",
    )

    if uploaded_file is not None:
        raw_text, tables = extract_docx_content(uploaded_file)
        extracted_vocab = extract_vocab_pairs_from_tables(tables)
        extracted_exercises = split_into_exercises(raw_text)
        fill_blank_items = build_fill_blank_items(extracted_exercises, raw_text)
        exercise_quiz_items = build_exercise_quiz_items(extracted_exercises, raw_text)

        default_title = Path(uploaded_file.name).stem
        lesson_title = st.text_input("Lesson title", value=default_title, key="upload_lesson_title")

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Tables found", len(tables))
        c2.metric("Vocab pairs found", len(extracted_vocab))
        c3.metric("Exercises found", len(extracted_exercises))
        c4.metric("Answer-ready blanks", len(exercise_quiz_items))

        with st.expander("Preview extracted vocabulary", expanded=True):
            if extracted_vocab:
                st.dataframe(
                    pd.DataFrame(extracted_vocab)[["term", "meaning", "source"]],
                    use_container_width=True,
                    hide_index=True,
                )
            else:
                st.info("No vocab table pairs were detected in this file.")

        with st.expander("Preview extracted exercises", expanded=True):
            if extracted_exercises:
                preview_rows = [
                    {
                        "position": item["position"],
                        "type": item["exercise_type"],
                        "prompt": item["prompt"],
                        "answer": item["answer"],
                    }
                    for item in extracted_exercises
                ]
                st.dataframe(pd.DataFrame(preview_rows), use_container_width=True, hide_index=True)
            else:
                st.info("No exercises were detected from the lesson text.")

        with st.expander("Preview lesson text"):
            if raw_text.strip():
                st.text_area("Raw lesson text", raw_text, height=280)
            else:
                st.warning("No paragraph text was extracted from the document.")

        if st.button("Save lesson", type="primary", key="save_lesson_button"):
            try:
                lesson_id, lesson_count = save_lesson_record(
                    title=lesson_title,
                    source_filename=uploaded_file.name,
                    raw_text=raw_text,
                )
                save_exercise_records(lesson_id, extracted_exercises)
                save_vocab_records(lesson_id, extracted_vocab)

                st.session_state["last_save_message"] = (
                    f"Saved Lesson #{lesson_count} with "
                    f"{len(extracted_exercises)} exercises and {len(extracted_vocab)} vocab items."
                )
                st.rerun()
            except Exception as e:
                st.error(f"Save failed: {e}")


# -------------------------------------------------
# Tab 2: Lesson History
# -------------------------------------------------
with tabs[1]:
    st.subheader("Lesson history")

    if not lessons:
        st.info("No lessons saved yet. Upload a .docx lesson first.")
    else:
        selected_label = st.selectbox(
            "Choose a saved lesson",
            options=list(lesson_map.keys()),
            key="history_selected_lesson",
        )
        selected_lesson_id = int(lesson_map[selected_label])

        lesson = fetch_lesson_by_id(selected_lesson_id)
        vocab_items = fetch_vocab_for_lesson(selected_lesson_id)
        exercises = fetch_exercises_for_lesson(selected_lesson_id)

        meta1, meta2, meta3 = st.columns(3)
        meta1.metric("Vocabulary items", len(vocab_items))
        meta2.metric("Exercises", len(exercises))
        meta3.metric("Created", lesson["created_at"] if lesson else "—")

        if lesson:
            st.markdown(f"### Lesson {display_numbers[selected_lesson_id]} — {lesson['title']}")
            if lesson["source_filename"]:
                st.caption(f"Source file: {lesson['source_filename']}")

        action_col1, action_col2 = st.columns(2)
        if action_col1.button("Delete selected lesson", key="delete_selected_lesson_button"):
            delete_lesson_record(selected_lesson_id)
            st.success("Selected lesson deleted.")
            st.rerun()

        if action_col2.button("Delete last 2 lessons", key="delete_last_two_lessons_button"):
            deleted_count = delete_last_n_lessons(2)
            st.success(f"Deleted {deleted_count} lesson(s).")
            st.rerun()

        st.markdown("---")

        with st.expander("Vocabulary", expanded=True):
            if vocab_items:
                st.dataframe(
                    pd.DataFrame(vocab_items)[["term", "meaning", "source", "position"]],
                    use_container_width=True,
                    hide_index=True,
                )
            else:
                st.info("No vocabulary saved for this lesson.")

        with st.expander("Exercises", expanded=True):
            if exercises:
                exercise_rows = [
                    {
                        "position": item["position"],
                        "title": item["exercise_title"],
                        "type": item["exercise_type"],
                        "prompt": item["prompt"],
                        "answer": item["answer"],
                    }
                    for item in exercises
                ]
                st.dataframe(pd.DataFrame(exercise_rows), use_container_width=True, hide_index=True)
            else:
                st.info("No exercises saved for this lesson.")

        with st.expander("Lesson text"):
            if lesson and lesson["raw_text"]:
                st.text_area("Saved text", lesson["raw_text"], height=300)
            else:
                st.info("No raw lesson text saved.")


# -------------------------------------------------
# Tab 3: Flashcards
# -------------------------------------------------
with tabs[2]:
    st.subheader("Flashcards")

    if not lessons:
        st.info("No lessons available yet.")
    else:
        selected_label = st.selectbox(
            "Choose a lesson for flashcards",
            options=list(lesson_map.keys()),
            key="flashcards_selected_lesson",
        )
        selected_lesson_id = int(lesson_map[selected_label])
        vocab_items = fetch_vocab_for_lesson(selected_lesson_id)

        if not vocab_items:
            st.warning("This lesson has no vocabulary pairs saved yet.")
        else:
            col1, col2, col3 = st.columns([2, 1, 1])

            order_mode = col1.selectbox(
                "Card order",
                options=["A → Z", "Random starting point", "Full shuffle", "Start from letter"],
                key="flashcard_order_mode",
            )

            start_letter = ""
            if order_mode == "Start from letter":
                start_letter = col2.text_input("Starting letter", value="A", max_chars=1, key="flashcard_start_letter")
            else:
                col2.write("")
                col2.write("")

            reset_clicked = col3.button("Reset deck", key="reset_flashcards_button")

            current_signature = flashcard_signature(selected_lesson_id, order_mode, start_letter, vocab_items)
            if reset_clicked or st.session_state.flashcard_signature != current_signature:
                reset_flashcards(selected_lesson_id, vocab_items, order_mode, start_letter)

            deck = st.session_state.flashcard_deck
            if not deck:
                st.info("No flashcards available.")
            else:
                idx = max(0, min(st.session_state.flashcard_index, len(deck) - 1))
                st.session_state.flashcard_index = idx
                card = deck[idx]

                top1, top2, top3 = st.columns(3)
                top1.metric("Card", f"{idx + 1} / {len(deck)}")
                top2.metric("Term", card["term"])
                top3.metric("Mode", order_mode)

                st.markdown('<div class="section-label">Front</div>', unsafe_allow_html=True)
                render_big_card(card["term"])

                st.markdown('<div class="section-label">Back</div>', unsafe_allow_html=True)
                if st.session_state.flashcard_show_back:
                    render_big_card(card["meaning"] or "No meaning saved for this card.", answer=True)
                else:
                    st.caption("Click **Flip card** to reveal the meaning.")

                nav1, nav2, nav3, nav4 = st.columns(4)

                if nav1.button("Previous", key="flashcard_prev"):
                    st.session_state.flashcard_index = max(0, st.session_state.flashcard_index - 1)
                    st.session_state.flashcard_show_back = False
                    st.rerun()

                if nav2.button("Flip card", key="flashcard_flip"):
                    st.session_state.flashcard_show_back = not st.session_state.flashcard_show_back
                    st.rerun()

                if nav3.button("Next", key="flashcard_next"):
                    st.session_state.flashcard_index = min(len(deck) - 1, st.session_state.flashcard_index + 1)
                    st.session_state.flashcard_show_back = False
                    st.rerun()

                if nav4.button("Restart", key="flashcard_restart"):
                    reset_flashcards(selected_lesson_id, vocab_items, order_mode, start_letter)
                    st.rerun()


# -------------------------------------------------
# Tab 4: Quiz Mode
# -------------------------------------------------
with tabs[3]:
    st.subheader("Quiz Mode")

    if not lessons:
        st.info("No lessons available yet.")
    else:
        selected_label = st.selectbox(
            "Choose a lesson for quiz mode",
            options=list(lesson_map.keys()),
            key="quiz_selected_lesson",
        )
        selected_lesson_id = int(lesson_map[selected_label])

        lesson = fetch_lesson_by_id(selected_lesson_id)
        vocab_items = fetch_vocab_for_lesson(selected_lesson_id)
        exercises = fetch_exercises_for_lesson(selected_lesson_id)
        lesson_text = lesson["raw_text"] if lesson else ""
        exercise_quiz_items = build_exercise_quiz_items(exercises, lesson_text)

        available_quiz_types: List[str] = []
        if vocab_items:
            available_quiz_types.append("Vocabulary Quiz")
        if exercise_quiz_items:
            available_quiz_types.append("Exercise Quiz")

        if not available_quiz_types:
            st.warning("This lesson has no vocabulary pairs or answer-ready blank exercises yet.")
        else:
            quiz_type = st.radio(
                "Quiz type",
                options=available_quiz_types,
                horizontal=True,
                key="quiz_type_selector",
            )

            if quiz_type == "Vocabulary Quiz":
                direction = st.radio(
                    "Quiz direction",
                    options=["French → English", "English → French"],
                    horizontal=True,
                    key="quiz_direction",
                )

                current_signature = quiz_signature(selected_lesson_id, direction, vocab_items)
                if st.session_state.quiz_signature != current_signature:
                    reset_quiz(selected_lesson_id, vocab_items, direction)

                deck = st.session_state.quiz_deck
                if not deck:
                    st.info("No vocabulary quiz items available.")
                else:
                    idx = max(0, min(st.session_state.quiz_index, len(deck) - 1))
                    st.session_state.quiz_index = idx
                    card = deck[idx]

                    prompt = card["term"] if direction == "French → English" else (card["meaning"] or card["term"])
                    expected = card["meaning"] if direction == "French → English" else card["term"]

                    q1, q2, q3 = st.columns(3)
                    q1.metric("Question", f"{idx + 1} / {len(deck)}")
                    q2.metric("Score", st.session_state.quiz_score)
                    q3.metric("Type", "Vocabulary")

                    st.markdown('<div class="section-label">Translate this</div>', unsafe_allow_html=True)
                    render_big_card(prompt)

                    user_answer = st.text_input("Your answer", key=f"quiz_input_box_{st.session_state.get('quiz_input_nonce', 0)}")
                    st.session_state.quiz_input = user_answer

                    b1, b2, b3, b4 = st.columns(4)
                    if b1.button("Check answer", key="quiz_check"):
                        if not st.session_state.quiz_answered:
                            if is_answer_correct(user_answer, expected):
                                st.session_state.quiz_feedback = f"✅ Correct. Answer: {expected}"
                                st.session_state.quiz_score += 1
                            else:
                                st.session_state.quiz_feedback = f"❌ Not quite. Correct answer: {expected}"
                            st.session_state.quiz_answered = True
                            st.rerun()

                    if b2.button("Show answer", key="quiz_show_answer"):
                        if not st.session_state.quiz_answered:
                            st.session_state.quiz_feedback = f"📘 Answer: {expected}"
                            st.session_state.quiz_answered = True
                            st.rerun()

                    if b3.button("Next question", key="quiz_next"):
                        if st.session_state.quiz_index < len(deck) - 1:
                            st.session_state.quiz_index += 1
                        st.session_state.quiz_feedback = ""
                        st.session_state.quiz_answered = False
                        st.session_state.quiz_input = ""
                        st.session_state.quiz_input_nonce = st.session_state.get("quiz_input_nonce", 0) + 1
                        st.rerun()

                    if b4.button("Restart quiz", key="quiz_restart"):
                        reset_quiz(selected_lesson_id, vocab_items, direction)
                        st.rerun()

                    if st.session_state.quiz_feedback:
                        st.write(st.session_state.quiz_feedback)

            else:
                current_signature = exercise_quiz_signature(selected_lesson_id, exercise_quiz_items)
                if st.session_state.exercise_quiz_signature != current_signature:
                    reset_exercise_quiz(selected_lesson_id, exercise_quiz_items)

                deck = st.session_state.exercise_quiz_deck
                if not deck:
                    st.info("No exercise quiz items available.")
                else:
                    idx = max(0, min(st.session_state.exercise_quiz_index, len(deck) - 1))
                    st.session_state.exercise_quiz_index = idx
                    item = deck[idx]

                    q1, q2, q3 = st.columns(3)
                    q1.metric("Question", f"{idx + 1} / {len(deck)}")
                    q2.metric("Score", st.session_state.exercise_quiz_score)
                    q3.metric("Type", "Fill-in-the-blank")

                    st.markdown('<div class="section-label">Complete this</div>', unsafe_allow_html=True)
                    render_big_card(item["prompt"])

                    user_answer = st.text_input(
                        "Your answer",
                        key=f"exercise_quiz_input_box_{st.session_state.get('exercise_quiz_input_nonce', 0)}",
                    )
                    st.session_state.exercise_quiz_input = user_answer

                    b1, b2, b3, b4 = st.columns(4)
                    if b1.button("Check answer", key="exercise_quiz_check"):
                        if not st.session_state.exercise_quiz_answered:
                            if is_answer_correct(user_answer, item["answer"]):
                                st.session_state.exercise_quiz_feedback = f"✅ Correct. Answer: {item['answer']}"
                                st.session_state.exercise_quiz_score += 1
                            else:
                                st.session_state.exercise_quiz_feedback = f"❌ Not quite. Correct answer: {item['answer']}"
                            st.session_state.exercise_quiz_answered = True
                            st.rerun()

                    if b2.button("Show answer", key="exercise_quiz_show_answer"):
                        if not st.session_state.exercise_quiz_answered:
                            st.session_state.exercise_quiz_feedback = f"📘 Answer: {item['answer']}"
                            st.session_state.exercise_quiz_answered = True
                            st.rerun()

                    if b3.button("Next question", key="exercise_quiz_next"):
                        if st.session_state.exercise_quiz_index < len(deck) - 1:
                            st.session_state.exercise_quiz_index += 1
                        st.session_state.exercise_quiz_feedback = ""
                        st.session_state.exercise_quiz_answered = False
                        st.session_state.exercise_quiz_input = ""
                        st.session_state.exercise_quiz_input_nonce = st.session_state.get("exercise_quiz_input_nonce", 0) + 1
                        st.rerun()

                    if b4.button("Restart quiz", key="exercise_quiz_restart"):
                        reset_exercise_quiz(selected_lesson_id, exercise_quiz_items)
                        st.rerun()

                    if st.session_state.exercise_quiz_feedback:
                        st.write(st.session_state.exercise_quiz_feedback)


# -------------------------------------------------
# Tab 5: Exercise Practice
# -------------------------------------------------
with tabs[4]:
    st.subheader("Fill-in-the-blank Practice")

    if not lessons:
        st.info("No lessons available yet.")
    else:
        selected_label = st.selectbox(
            "Choose a lesson for exercise practice",
            options=list(lesson_map.keys()),
            key="practice_selected_lesson",
        )
        selected_lesson_id = int(lesson_map[selected_label])

        lesson = fetch_lesson_by_id(selected_lesson_id)
        exercises = fetch_exercises_for_lesson(selected_lesson_id)
        vocab_items = fetch_vocab_for_lesson(selected_lesson_id)
        lesson_text = lesson["raw_text"] if lesson else ""

        practice_items = build_fill_blank_items(exercises, lesson_text)

        if not practice_items:
            st.warning("No blank-style prompts were found in this lesson.")
        else:
            show_hints = st.checkbox("Show vocab hints", value=True, key="practice_show_hints")

            current_signature = practice_signature(selected_lesson_id, practice_items)
            if st.session_state.practice_signature != current_signature:
                reset_practice(selected_lesson_id, practice_items)

            deck = st.session_state.practice_deck
            idx = max(0, min(st.session_state.practice_index, len(deck) - 1))
            st.session_state.practice_index = idx
            item = deck[idx]

            top1, top2 = st.columns(2)
            top1.metric("Prompt", f"{idx + 1} / {len(deck)}")
            top2.metric("Has stored answer", "Yes" if clean_text(item["answer"]) else "No")

            st.markdown('<div class="section-label">Prompt</div>', unsafe_allow_html=True)
            render_big_card(item["prompt"])

            practice_answer = st.text_input(
                "Your answer",
                key=f"practice_input_box_{st.session_state.get('practice_input_nonce', 0)}",
            )
            st.session_state.practice_input = practice_answer

            if show_hints:
                with st.expander("Vocab hints"):
                    hint_pool = build_vocab_hint_pool(vocab_items, limit=10)
                    if hint_pool:
                        st.dataframe(
                            pd.DataFrame(hint_pool)[["term", "meaning"]],
                            use_container_width=True,
                            hide_index=True,
                        )
                    else:
                        st.info("No vocabulary hints saved for this lesson.")

            b1, b2, b3, b4 = st.columns(4)
            if b1.button("Check", key="practice_check"):
                expected = clean_text(item["answer"])
                if expected:
                    if is_answer_correct(practice_answer, expected):
                        st.session_state.practice_feedback = f"✅ Correct. Answer: {expected}"
                    else:
                        st.session_state.practice_feedback = f"❌ Not quite. Correct answer: {expected}"
                else:
                    st.session_state.practice_feedback = "ℹ️ No stored answer for this blank yet. Use the hints and review manually."
                st.session_state.practice_answered = True
                st.rerun()

            if b2.button("Reveal answer", key="practice_reveal"):
                expected = clean_text(item["answer"])
                if expected:
                    st.session_state.practice_feedback = f"📘 Answer: {expected}"
                else:
                    st.session_state.practice_feedback = "ℹ️ No stored answer is saved for this prompt yet."
                st.session_state.practice_answered = True
                st.rerun()

            if b3.button("Next", key="practice_next"):
                if st.session_state.practice_index < len(deck) - 1:
                    st.session_state.practice_index += 1
                st.session_state.practice_feedback = ""
                st.session_state.practice_answered = False
                st.session_state.practice_input = ""
                st.session_state.practice_input_nonce = st.session_state.get("practice_input_nonce", 0) + 1
                st.rerun()

            if b4.button("Restart practice", key="practice_restart"):
                reset_practice(selected_lesson_id, practice_items)
                st.rerun()

            if st.session_state.practice_feedback:
                st.write(st.session_state.practice_feedback)